"""Word(.docx) 형식으로 업무보고 취합본 생성. 한글 글꼴 네이티브 지원.

원본 HWPX(04.24) 레이아웃 복제:
  Page 1: 사업단 공통확인사항 1 (정지수)
  Page 2: 사업단 공통확인사항 2 (최혜민, 실적/계획 2칸)
  Pages 3-12: 팀원 10명 1인 1페이지 (6열: 과제|분야|이름|구분|실적|계획)
  Page 13: 스마트돌봄스페이스 + 회의자료 3종
"""
import io
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from team_config import TEAM_MEMBERS

FONT_NAME = "맑은 고딕"  # Windows 기본. Word가 fallback 잘 처리
BLUE = RGBColor(0x00, 0x00, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY_BG = "F5F5F5"  # 헤더 배경색


def _set_cell_vertical_text(cell):
    """셀 텍스트를 세로쓰기로 (한 글자씩 위→아래)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    text_direction = OxmlElement('w:textDirection')
    text_direction.set(qn('w:val'), 'btLr')  # bottom-to-top, left-to-right = 세로쓰기
    tc_pr.append(text_direction)


def _set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell):
    """테두리 설정 (모든 변)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')  # 0.5pt
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tc_pr.append(borders)


def _add_text(cell, text: str, bold=False, color=None, align_center=False,
              font_size=9):
    """셀에 텍스트 추가 (기존 내용 유지하며 덧붙임)."""
    p = cell.paragraphs[-1] if cell.text == "" else cell.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate((text or "").split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = FONT_NAME
        r.font.size = Pt(font_size)
        if bold:
            r.bold = True
        if color is not None:
            r.font.color.rgb = color
        # 한글 폰트 강제 지정
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)


def _set_cell_text(cell, text: str, bold=False, color=None,
                   align_center=False, font_size=9, vertical=False,
                   bg=None):
    """셀 내용을 텍스트로 설정 (초기화 후 채움). 세로쓰기/배경색 옵션."""
    # 기존 paragraph 첫 번째만 쓰고 나머지 삭제
    for p in list(cell.paragraphs):
        p_el = p._element
        p_el.getparent().remove(p_el)
    # 새 paragraph 추가
    p = cell.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate((text or "").split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = FONT_NAME
        r.font.size = Pt(font_size)
        if bold:
            r.bold = True
        if color is not None:
            r.font.color.rgb = color
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if align_center else WD_ALIGN_VERTICAL.TOP
    if vertical:
        _set_cell_vertical_text(cell)
    if bg:
        _set_cell_bg(cell, bg)


def _add_title(doc, title_date: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"과업별 업무 보고 ({title_date})")
    r.font.name = FONT_NAME
    r.font.size = Pt(14)
    r.bold = True
    r.underline = True
    rPr = r._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rPr.append(rFonts)


def _landscape_section(doc, first=False):
    """가로 A4 섹션 추가."""
    if first:
        section = doc.sections[0]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(10)


def build_docx(submissions: dict, title_date: str,
               period_start: str, period_end: str,
               plan_start: str, plan_end: str) -> bytes:
    doc = Document()

    period_header = f"업무 실적({period_start} ~ {period_end})"
    plan_header = f"업무 계획({plan_start} ~ {plan_end})"

    # ─── Page 1: 사업단 공통확인사항 1 ───
    _landscape_section(doc, first=True)
    _add_title(doc, title_date)

    jjs = submissions.get("정지수", {})
    tbl = doc.add_table(rows=1, cols=2)
    tbl.columns[0].width = Mm(40)
    tbl.columns[1].width = Mm(237)
    row = tbl.rows[0]
    row.height = Mm(150)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_cell_text(row.cells[0], "사업단 공통\n확인사항 1",
                   bold=True, align_center=True, font_size=11,
                   vertical=True, bg=GREY_BG)
    _set_cell_text(row.cells[1],
                   jjs.get("project_confirmation_1", "") or "(미작성)",
                   font_size=10)
    for c in row.cells:
        _set_cell_border(c)

    # ─── Page 2: 사업단 공통확인사항 2 ───
    _landscape_section(doc)
    _add_title(doc, title_date)

    chm = submissions.get("최혜민", {})
    tbl = doc.add_table(rows=2, cols=3)
    tbl.columns[0].width = Mm(40)
    tbl.columns[1].width = Mm(118)
    tbl.columns[2].width = Mm(119)

    # 헤더 행
    header_row = tbl.rows[0]
    header_row.height = Mm(10)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_cell_text(header_row.cells[1], period_header,
                   bold=True, align_center=True, font_size=10, bg=GREY_BG)
    _set_cell_text(header_row.cells[2], plan_header,
                   bold=True, align_center=True, font_size=10, bg=GREY_BG)
    # 본문 행
    body_row = tbl.rows[1]
    body_row.height = Mm(160)
    body_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_cell_text(body_row.cells[1],
                   chm.get("project_confirmation_2_done", "") or "(미작성)",
                   font_size=10)
    _set_cell_text(body_row.cells[2],
                   chm.get("project_confirmation_2_plan", "") or "(미작성)",
                   font_size=10)
    # 첫 컬럼 병합
    merged = header_row.cells[0].merge(body_row.cells[0])
    _set_cell_text(merged, "사업단 공통\n확인사항 2",
                   bold=True, align_center=True, font_size=11,
                   vertical=True, bg=GREY_BG)
    for r in tbl.rows:
        for c in r.cells:
            _set_cell_border(c)

    # ─── Pages 3-12: 팀원 1명씩 ───
    for m in TEAM_MEMBERS:
        _landscape_section(doc)
        _add_title(doc, title_date)
        _build_member_table(doc, m, submissions.get(m["name"], {}),
                            period_header, plan_header)

    # ─── 마지막 페이지: 스마트돌봄스페이스 + 회의자료 3종 ───
    _landscape_section(doc)
    _add_title(doc, title_date)
    _build_bottom_table(doc, submissions, period_header, plan_header)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_member_table(doc, member, data, period_header, plan_header):
    """팀원 1명의 한 페이지짜리 표 (6열 구조: 과제|분야|이름|구분|실적|계획)."""
    cat1 = member["category1"]
    cat2 = member["category2"]
    name = member["name"]

    if member["has_research"]:
        # 4행: [header][획득][연구 실적/계획][업무 실적/계획]
        tbl = doc.add_table(rows=4, cols=6)
        for col, w in zip(tbl.columns, [12, 14, 14, 12, 115, 110]):
            col.width = Mm(w)
        row_heights = [8, 15, 60, 75]

        # 헤더 행
        hdr = tbl.rows[0]
        # 구분 4열 병합
        m01 = hdr.cells[0].merge(hdr.cells[1]).merge(hdr.cells[2]).merge(hdr.cells[3])
        _set_cell_text(m01, "구분", bold=True, align_center=True,
                       font_size=10, bg=GREY_BG)
        _set_cell_text(hdr.cells[4], period_header,
                       bold=True, align_center=True, font_size=10, bg=GREY_BG)
        _set_cell_text(hdr.cells[5], plan_header,
                       bold=True, align_center=True, font_size=10, bg=GREY_BG)

        # 획득데이터 행 (r=1)
        acq_row = tbl.rows[1]
        _set_cell_text(acq_row.cells[0], cat1, bold=True, align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(acq_row.cells[1], cat2, align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(acq_row.cells[2], name, bold=True, align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(acq_row.cells[3], "연구", align_center=True,
                       vertical=True, font_size=10)
        # c4-c5 병합 획득데이터 (먼저 병합한 후 내용 입력 — 중복 방지)
        acq_merged = acq_row.cells[4].merge(acq_row.cells[5])
        acq_text = data.get("acquired_data", "").strip()
        # 이미 "획득 데이터:" 로 시작하면 그대로, 아니면 prefix 추가
        if acq_text and not acq_text.startswith("획득 데이터"):
            display_text = f"획득 데이터: {acq_text}"
        else:
            display_text = acq_text or "획득 데이터:"
        _set_cell_text(acq_merged, display_text,
                       bold=True, color=BLUE, font_size=9)

        # 연구 실적/계획 행 (r=2)
        r_row = tbl.rows[2]
        _set_cell_text(r_row.cells[4], data.get("research_done", ""),
                       font_size=9)
        _set_cell_text(r_row.cells[5], data.get("research_plan", ""),
                       font_size=9)

        # 업무 행 (r=3)
        t_row = tbl.rows[3]
        _set_cell_text(t_row.cells[3], "업무", align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(t_row.cells[4], data.get("task_done", ""), font_size=9)
        _set_cell_text(t_row.cells[5], data.get("task_plan", ""), font_size=9)

        # 0-2열을 r1~r3 병합 (과제/분야/이름 3행 세로 병합)
        for col_idx in [0, 1, 2]:
            merged = tbl.rows[1].cells[col_idx]
            for row_idx in [2, 3]:
                merged = merged.merge(tbl.rows[row_idx].cells[col_idx])

        # 연구 셀 (col 3) r1~r2 병합
        research_merged = tbl.rows[1].cells[3].merge(tbl.rows[2].cells[3])

        # 행 높이
        for r, h in zip(tbl.rows, row_heights):
            r.height = Mm(h)
            r.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    else:
        # 2행: 헤더 + 업무 행
        tbl = doc.add_table(rows=2, cols=6)
        for col, w in zip(tbl.columns, [12, 14, 14, 12, 115, 110]):
            col.width = Mm(w)

        hdr = tbl.rows[0]
        m01 = hdr.cells[0].merge(hdr.cells[1]).merge(hdr.cells[2]).merge(hdr.cells[3])
        _set_cell_text(m01, "구분", bold=True, align_center=True,
                       font_size=10, bg=GREY_BG)
        _set_cell_text(hdr.cells[4], period_header,
                       bold=True, align_center=True, font_size=10, bg=GREY_BG)
        _set_cell_text(hdr.cells[5], plan_header,
                       bold=True, align_center=True, font_size=10, bg=GREY_BG)

        t_row = tbl.rows[1]
        _set_cell_text(t_row.cells[0], cat1, bold=True, align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(t_row.cells[1], cat2, align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(t_row.cells[2], name, bold=True, align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(t_row.cells[3], "업무", align_center=True,
                       vertical=True, font_size=10)
        _set_cell_text(t_row.cells[4], data.get("task_done", ""), font_size=9)
        _set_cell_text(t_row.cells[5], data.get("task_plan", ""), font_size=9)

        tbl.rows[0].height = Mm(8)
        tbl.rows[1].height = Mm(150)
        for r in tbl.rows:
            r.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for r in tbl.rows:
        for c in r.cells:
            _set_cell_border(c)


def _build_bottom_table(doc, submissions, period_header, plan_header):
    """스마트돌봄스페이스 + 회의자료 3종."""
    bjs = submissions.get("백정은", {})
    chm = submissions.get("최혜민", {})

    tbl = doc.add_table(rows=5, cols=3)
    tbl.columns[0].width = Mm(40)
    tbl.columns[1].width = Mm(118)
    tbl.columns[2].width = Mm(119)

    hdr = tbl.rows[0]
    _set_cell_text(hdr.cells[0], "구분", bold=True, align_center=True,
                   font_size=10, bg=GREY_BG)
    _set_cell_text(hdr.cells[1], period_header, bold=True, align_center=True,
                   font_size=10, bg=GREY_BG)
    _set_cell_text(hdr.cells[2], plan_header, bold=True, align_center=True,
                   font_size=10, bg=GREY_BG)

    sections = [
        ("스마트돌봄\n스페이스",
         bjs.get("smart_care_space_done", ""),
         bjs.get("smart_care_space_plan", "")),
        ("1. 연구소\n회의자료\n(소장주재회의)",
         chm.get("research_meeting", ""), ""),
        ("2. 원장+재활원\n주요간부\n회의자료\n(주간 현안보고)",
         chm.get("director_meeting", ""), ""),
        ("3. 복지부 본부\n주간일정\n보산진 보고\n(의료기기R&D)",
         chm.get("mohw_weekly", ""), ""),
    ]
    row_heights = [10, 30, 30, 30, 50]

    for idx, (label, done, plan) in enumerate(sections, start=1):
        row = tbl.rows[idx]
        _set_cell_text(row.cells[0], label, bold=True, align_center=True,
                       font_size=9, bg=GREY_BG)
        _set_cell_text(row.cells[1], done, font_size=9)
        _set_cell_text(row.cells[2], plan, font_size=9)

    for r, h in zip(tbl.rows, row_heights):
        r.height = Mm(h)
        r.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for r in tbl.rows:
        for c in r.cells:
            _set_cell_border(c)
